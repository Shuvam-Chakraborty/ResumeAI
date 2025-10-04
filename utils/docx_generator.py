import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_docx(resume_data: dict) -> bytes:
    """Generate DOCX from resume data"""
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Header - Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(resume_data.get('name', ''))
    name_run.font.size = Pt(18)
    name_run.bold = True
    
    # Contact Info
    contact_parts = []
    for field in ['location', 'email', 'phone', 'linkedin', 'github']:
        if resume_data.get(field):
            contact_parts.append(resume_data[field])
    
    if contact_parts:
        contact_para = doc.add_paragraph(' • '.join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Professional Summary
    if resume_data.get('summary'):
        doc.add_heading('PROFESSIONAL SUMMARY', level=2)
        doc.add_paragraph(resume_data['summary'])
    
    # Experience
    if resume_data.get('experiences'):
        doc.add_heading('PROFESSIONAL EXPERIENCE', level=2)
        for exp in resume_data['experiences']:
            p = doc.add_paragraph()
            p.add_run(exp.get('role', '')).bold = True
            
            company_text = exp.get('company', '')
            if exp.get('location'):
                company_text += f" • {exp['location']}"
            company_text += f" • {exp.get('duration', '')}"
            doc.add_paragraph(company_text)
            
            for detail in exp.get('details', []):
                doc.add_paragraph(detail, style='List Bullet')
            
            doc.add_paragraph()
    
    # Projects
    if resume_data.get('projects'):
        doc.add_heading('PROJECTS', level=2)
        for proj in resume_data['projects']:
            p = doc.add_paragraph()
            p.add_run(proj.get('name', '')).bold = True
            
            tech_text = f"{proj.get('technologies', '')} • {proj.get('duration', '')}"
            doc.add_paragraph(tech_text)
            
            for detail in proj.get('details', []):
                doc.add_paragraph(detail, style='List Bullet')
            
            doc.add_paragraph()
    
    # Education
    if resume_data.get('education'):
        doc.add_heading('EDUCATION', level=2)
        for edu in resume_data['education']:
            p = doc.add_paragraph()
            p.add_run(edu.get('degree', '')).bold = True
            
            school_text = edu.get('school', '')
            if edu.get('details'):
                school_text += f" • {edu['details']}"
            school_text += f" • {edu.get('year', '')}"
            doc.add_paragraph(school_text)
    
    # Skills
    if resume_data.get('skills'):
        doc.add_heading('SKILLS', level=2)
        doc.add_paragraph(resume_data['skills'])
    
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()